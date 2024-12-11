#pylint: disable=C0303
#pylint: disable=C0301
#pylint: disable=W0718

"""Importing all the necessary libraries"""

import os
import io
import json
import openai
import PyPDF2
from docx import Document
import requests
from pathlib import Path
from typing import Dict, Any
from datetime import datetime, timedelta
from app.core.logger_setup import setup_logger
from azure.storage.blob import BlobServiceClient, BlobSasPermissions, generate_blob_sas
from app.core.config import settings

# Setup Logger
logger = setup_logger(__name__)

# Configure OpenAI settings
openai.api_key = settings.OPENAI_API_KEY
openai.api_base = settings.OPENAI_API_BASE
openai.api_version = settings.OPENAI_API_VERSION
openai.api_type = settings.OPENAI_API_TYPE

# Configure Azure Blob Storage settings
account_name = settings.ACCOUNT_NAME
container_name = settings.CONTAINER_NAME
account_key = settings.ACCOUNT_KEY
account_url = settings.BLOB_ACCOUNT_URL
deployment_name = settings.OPENAI_DEPLOYMENT_NAME

class CandidateJobEvaluator:
    """
        Handles evaluation and processing of candidate resumes and job descriptions.
        Provides functionality for resume parsing, job description formatting,
        and similarity scoring between resumes and job descriptions.
    """
    def __init__(self):
        self.blob_service_client = BlobServiceClient(account_url=account_url, credential=account_key)
        self.formatted_resume = None
        logger.info(f"BlobServiceClient initialized for account: {account_name}")

    def generate_sas_token(self, container_name: str, blob_name: str) -> str:
        """
        Generate a SAS (Shared Access Signature) token for a specific blob.
        """
        try:
            sas_token = generate_blob_sas(
                account_name=account_name,
                container_name=container_name,
                blob_name=blob_name,
                account_key=account_key,
                permission=BlobSasPermissions(read=True, write=True),
                expiry=datetime.utcnow() + timedelta(hours=0.166667)
            )
            logger.info(f"SAS token generated successfully: {sas_token}")
            return sas_token
        except Exception as e:
            logger.error(f"Error generating SAS token: {e}")
            raise

    def upload_pdf_to_blob(self, pdf_filename: str) -> bool:
        """
        Upload a PDF file to Azure Blob Storage.
        """
        try:
            pdf_path = settings.RESUMES_DIR / pdf_filename
            logger.info(f"Looking for file at path: {pdf_path}")

            if not os.path.exists(pdf_path):
                logger.error(f"File not found: {pdf_path}")
                return False

            sas_token = self.generate_sas_token(container_name, pdf_filename)
            blob_url = f"{account_url}/{container_name}/{pdf_filename}?{sas_token}"
            blob_pdf_path = f"{account_url}/{container_name}/{pdf_filename}"

            with open(pdf_path, "rb") as pdf_file:
                data = pdf_file.read()
                headers = {
                    "x-ms-blob-type": "BlockBlob",
                    "Content-Type": "application/pdf"
                }
                response = requests.put(blob_url, data=data, headers=headers)

                if response.status_code == 201:
                    logger.info(f"Successfully uploaded blob: {pdf_filename}")
                    return True, blob_pdf_path
                else:
                    logger.error(f"Failed to upload blob: {pdf_filename}. Status code: {response.status_code}")
                    return False

        except Exception as e:
            logger.error(f"Error uploading blob: {str(e)}")
            return False

    def extract_text_from_pdf(self, file_content, filename):
        file_extension = filename.split('.')[-1]
        print('###########', file_extension)
        """Extract text from either PDF or DOCX file content."""
        print(f'Inside extract_text_from_document for {file_extension} file')
        text = ""
        
        try:
            if file_extension.lower() == '.pdf':
                # Create a PDF file object from bytes
                pdf_file = io.BytesIO(file_content)
                # Create PDF reader object
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                # Extract text from each page
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                    
            elif file_extension.lower() in ['docx', 'doc']:
                # Create a document object from bytes
                doc_file = io.BytesIO(file_content)
                doc = Document(doc_file)
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                    
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
                        text += "\n"  # Add extra newline between table rows
                    
            else:
                raise Exception(f"Unsupported file format: {file_extension}")
            
            return text.strip()
                
        except Exception as e:
            raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")
        
        return text

    def get_latest_blob_with_sas(self, filename) -> tuple[str, str, str]:
        """
        Get the latest blob from the container with its SAS token.
        """
        try:
            container_client = self.blob_service_client.get_container_client(container_name)
            blobs = list(container_client.list_blobs())

            if not blobs:
                logger.error("No blobs found in container")
                raise Exception("No blobs found in container")

            latest_blob = max(blobs, key=lambda x: x.last_modified)
            latest_blob_name = latest_blob.name

            logger.info(f"Latest blob found: {latest_blob_name}, Last modified: {latest_blob.last_modified}")

            sas_token = self.generate_sas_token(container_name, latest_blob_name)
            blob_url = f"{account_url}/{container_name}/{latest_blob_name}?{sas_token}"

            response = requests.get(blob_url)
            if response.status_code != 200:
                raise Exception(f"Failed to download blob: {response.status_code, filename}")

            # Extract text from PDF
            extracted_text = self.extract_text_from_pdf(response.content, filename)
            # print(f'get_latest_blob_with_sas:{extracted_text}')

            return latest_blob_name, blob_url, extracted_text

        except Exception as e:
            logger.error(f"Error getting latest blob with SAS: {str(e)}")
            raise

    def format_jd_with_gpt(self, jd_text):
        prompt = f"""
        You are a powerful AI that helps reformat job descriptions (JDs) into a structured JSON format. Below is the template you need to follow:

        **Template (in JSON format):**

        ```json
        {{
            "companyInfo": {{
                "companyName": "",
                "location": "",
                "industry": ""
            }},
            "requiredQualifications": [
                {{
                    "degree": "",
                    "field": "",
                    "additionalRequirements": ""
                }}
            ],
            "requiredWorkExperience": {{
                "yearsRequired": "",
                "description": ""
            }},
            "rolesAndResponsibilities": [
                ""
            ],
            "requiredSkills": [
                ""
            ],
            "requiredCertifications": [
                {{
                    "certificationName": "",
                    "issuingOrganization": ""
                }}
            ]
        }}
        ```

        **Job Description Text:**
        {jd_text}

        **Reformatted Job Description:**
        Please extract and reformat the information from the provided job description text according to the template above. Ensure that each section is clearly labeled and formatted appropriately. If any section is missing, leave it blank but indicate its presence in the template.
        """

        response = openai.ChatCompletion.create(
            engine=deployment_name,
            messages=[
                {"role": "system",
                 "content": prompt},
            ],
            max_tokens=1500,
            temperature=0.7
        )

        analysis = response.choices[0].message.content

        return analysis

    def format_resume_with_gpt(self, resume_text: str) -> Dict[str, Any]:
        print("inside format_resume_with_gpt")
        """
        Format resume text using GPT-4 into structured JSON
        """
        resume_text = resume_text.encode('utf-8', errors='ignore').decode('utf-8')
        
        # Remove the code block syntax and make it more direct
        system_message = """You are an AI that reformats resumes into structured JSON. 
        You must:
        1. Extract all relevant information
        2. Return ONLY valid JSON that can be parsed by json.loads()
        3. Do not include any markdown, code blocks, or additional text
        4. Ensure all text fields are properly escaped and on a single line
        5. Include all available skills, work experience, and education details"""

        prompt = f"""Format the following resume into valid JSON using this structure:
        {{
            "personalInfo": {{
                "name": "",
                "location": "",
                "email": "",
                "phone": "",
                "linkedIn": ""
            }},
            "education": [
                {{
                    "degree": "",
                    "field": "",
                    "institution": "",
                    "graduationYear": "",
                    "gpa": ""
                }}
            ],
            "workExperience": [
                {{
                    "companyName": "",
                    "position": "",
                    "duration": "",
                    "responsibilities": [""],
                    "achievements": [""]
                }}
            ],
            "skills": {{
                "technical": [""],
                "soft": [""],
                "languages": [""]
            }},
            "certifications": [
                {{
                    "name": "",
                    "issuingOrganization": "",
                    "issueDate": "",
                    "expiryDate": ""
                }}
            ],
            "projects": [
                {{
                    "name": "",
                    "description": "",
                    "technologies": [""],
                    "link": ""
                }}
            ]
        }}

        Resume text to format:
        {resume_text}

        IMPORTANT: Return only the JSON object itself, with no markdown formatting or code blocks."""

        try:
            response = openai.ChatCompletion.create(
                engine=deployment_name,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )

            response_text = response.choices[0].message.content.strip()
            
            # Clean up any potential markdown formatting
            if response_text.startswith('```'):
                response_text = response_text.split('```')[1]
                if response_text.startswith('json'):
                    response_text = response_text[4:]
            response_text = response_text.strip()

            try:
                parsed_json = json.loads(response_text)
                return parsed_json
            except json.JSONDecodeError as e:
                logger.error(f"Error parsing GPT response as JSON: {e}")
                logger.error(f"Raw response: {response_text}")
                return None

        except Exception as e:
            logger.error(f"Error in GPT API call: {e}")
            return None

    # def get_company_info_with_llm(self, company_name: str) -> Dict[str, Any]:
    #     """
    #     Use GPT to identify and provide information about a company.
        
    #     Args:
    #         company_name (str): Name of the company to search for
            
    #     Returns:
    #         Dict[str, Any]: Company information and possible LinkedIn URL
    #     """
    #     try:
    #         system_message = """You are a knowledgeable assistant that helps identify companies and provide accurate information about them.
    #         For any company name provided:
    #         1. Identify the most likely matching real company
    #         2. Provide key information about the company
    #         3. Consider searching for the company website and LinkedIn/Google presence
    #         4. If multiple companies match, provide information about the most relevant one
    #         5. If you're not confident about the company match, indicate that in your response
    #         6. For Indian companies, specifically consider:
    #         - Professional body registrations (ICAI, etc.)
    #         - Local business directories
    #         - Company websites (usually ending in .co.in or .com)
    #         Return ONLY valid JSON without any additional text."""

    #         prompt = f"""Please provide information about this company: {company_name}
            
    #         Think about how you would search for this company online:
    #         1. Common website patterns (company.com, company.co.in)
    #         2. Professional directories and associations
    #         3. Business registries
    #         4. Location-specific listings
            
    #         Return the information in this exact JSON format:
    #         {{
    #             "company_info": {{
    #                 "Website": "Company website URL if available, or most likely URL pattern",
    #                 "description": "A brief description covering the key points about the company",
    #                 "industry": "Main industry",
    #                 "Company size": "Approximate number of employees",
    #                 "Headquarters": "Main headquarters location",
    #                 "type": "Type of company (e.g., Private, Public, etc.)",
    #                 "key_products_services": ["List of main products/services"],
    #                 "Founded": "Founding year if known, otherwise null",
    #                 "Specialties": ["Any notable facts about the company"]
    #             }}
    #         }}"""

    #         try:
    #             response = openai.ChatCompletion.create(
    #                 engine=deployment_name,
    #                 messages=[
    #                     {"role": "system", "content": system_message},
    #                     {"role": "user", "content": prompt}
    #                 ],
    #                 max_tokens=2000,
    #                 temperature=0.7
    #             )

    #             response_text = response.choices[0].message.content.strip()

    #             try:
    #                 parsed_json = json.loads(response_text)
    #                 return parsed_json
    #             except json.JSONDecodeError as e:
    #                 logger.error(f"Error parsing GPT response as JSON: {e}")
    #                 logger.error(f"Raw response: {response_text}")
    #                 return None

    #         except Exception as e:
    #             logger.error(f"Error in GPT API call: {e}")
    #             return None

    #     except Exception as e:
    #         logger.error(f"Error getting company info with LLM: {str(e)}")
    #         return None

    def get_companies_info(self) -> Dict[str, Any]:
        """
        Get company background information matching the employment info format.
        """
        try:
            if not self.formatted_resume:
                raise ValueError("No formatted resume data available. Please process the resume first.")
                
            work_experience = self.formatted_resume.get("workExperience", [])
            if not work_experience:
                raise ValueError("No work experience found in the formatted resume")

            # Format company data like in the second code
            company_experience = []
            for exp in work_experience[:2]:  # Get two most recent companies
                exp_text = f"""
                Company: {exp.get('companyName')}
                Position: {exp.get('position')}
                Responsibilities: {', '.join(exp.get('responsibilities', []))}
                Achievements: {', '.join(exp.get('achievements', []))}
                """
                company_experience.append(exp_text)

            system_message = """You are a professional employment history summarizer. Create concise, 
            informative summaries of candidates' previous employment, focusing on the companies they worked for, 
            the companies' scope of business, and the candidate's role/function. Keep summaries to 2-3 sentences."""

            user_message = f"""
            Create a previous employment summary based on this information:

            Experience Details:
            {' '.join(company_experience)}

            Format the response like this example:
            Has worked in [Company Names] in past. These companies are involved in [business description] and 
            have [employee count] employees. This candidate worked in [specific part of the department] within 
            the [function/division] of the company.

            Important: Focus on factual information about previous employers and the candidate's roles."""

            response = openai.ChatCompletion.create(
                engine=deployment_name,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ],
                max_tokens=200,
                temperature=0.7
            )

            if response and response.choices:
                company_summary = response.choices[0].message.content.strip()
                return company_summary
            else:
                logger.error("No valid response from OpenAI API")
                return None

        except Exception as e:
            logger.error(f"Error in getting companies info: {str(e)}")
            return {"error": str(e)}

    def generate_similarity_scores(
        self,
        resume_text: Dict[str, Any],
        jd_data: Dict[str, Any],
        application_id: int
        ) -> Dict[str, Any]:
        try:
            logger.info(f"Generating similarity scores for application ID: {application_id}")

            prompt_template = """
            You are an AI assistant that calculates match percentages between resumes and job descriptions based on their JSON formats.
            Analyze the following JSONs and provide match percentages along with the relevant reasoning:
            # 1. Overall Relevance Score (0-100 points): Overall score based on entire evaluation done
            # 2. Skills Match (0-30 points): How well do the skills listed in the resume match the skills required in the JD?
            # 3. Experience Match (0-30 points): How well does the candidate's experience match the job requirements?
            # 4. Education Match (0-20 points): How well does the candidate's education align with the job requirements?
            # 5. Overall Relevance (0-20 points): Overall relevance of the candidate's profile to the job.

            Job Description: {processed_jd}
            Resume: {processed_resume}

            Calculate and provide these match percentages and give reasons for the match and not how the percentage is calculated:
 
            Skills Match: [Percentage]%
            (Calculate based on matching skills in resume["skills"] and resume["projects"]["technologiesUsed"] against jd["requiredSkills"])
 
            Experience Match: [Percentage]%
            (Calculate based on resume["workExperience"] match with jd["requiredWorkExperience"] considering years and responsibilities)
 
            Education Match: [Percentage]%
            (Calculate based on resume["qualifications"] and resume["certifications"] match with jd["requiredQualifications"] and jd["requiredCertifications"])
            
            Overall Match: [Average of all four percentages]%
            
            Overall Assessment:
             - Provide a brief evaluation of whether the candidate's experience aligns well with the role's requirements
            
            Potential Gaps:
            - List any critical responsibilities from the JD where the candidate's experience may be limited or lacking

            {{
            "matching_score": Overall Match Percentage,
            "sections": [
                {{
                "name": "Skills Match",
                "score": Skills Match Percentage,
                "max_score": 100,
                "overview": "Brief overview of skills match"
                }},
                {{
                "name": "Experience Match",
                "score": Experience Match Percentage,
                "max_score": 100,
                "overview": "Brief overview of experience match"
                }},
                {{
                "name": "Education Match",
                "score": Education Match Percentage,
                "max_score": 100,
                "overview": "Brief overview of education match"
                }},
                {{
                "name": "Overall Relevance",
                "score": Overall Match Percentage,
                "max_score": 100,
                "overview": "Overall assessment"
                }}
            ],
            "potential_gaps": [
                {{
                "description": "Potential Gaps"
                }}
            ]
            }}"""
            

            response = openai.ChatCompletion.create(
                engine=deployment_name,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant. Return ONLY the JSON object with no additional text."},
                    {"role": "user", "content": prompt_template.format(
                        processed_jd=json.dumps(jd_data, indent=2),
                        processed_resume=json.dumps(resume_text, indent=2)
                    )}
                ],
                max_tokens=3000,
                temperature=0.7
            )

            analysis = response.choices[0].message.content.strip()
            
            # Try to extract JSON from the response
            try:
                # First attempt: try to parse the entire response
                output = json.loads(analysis)
            except json.JSONDecodeError:
                # Second attempt: try to find and extract just the JSON object
                try:
                    # Find the first { and last }
                    start_idx = analysis.find('{')
                    end_idx = analysis.rstrip().rfind('}') + 1
                    if start_idx != -1 and end_idx != -1:
                        json_str = analysis[start_idx:end_idx]
                        output = json.loads(json_str)
                    else:
                        raise ValueError("Could not find valid JSON object in response")
                except Exception as e:
                    logger.error(f"Error extracting JSON from response: {str(e)}")
                    logger.error(f"Raw response: {analysis}")
                    raise

            return output

        except Exception as e:
            logger.error(f"Error generating similarity scores: {str(e)}")
            raise

    def process_resume(self, pdf_filename: str) -> dict:
        """
        Main method to handle the complete resume processing workflow.
        """
        try:
            # Step 1: Upload PDF to blob storage
            upload_success, blob_pdf_url = self.upload_pdf_to_blob(pdf_filename)
            if not upload_success:
                raise Exception("Failed to upload PDF to blob storage")

            latest_blob_name, blob_url, extracted_text = self.get_latest_blob_with_sas(pdf_filename)
            formatted_resume = self.format_resume_with_gpt(extracted_text)
            if formatted_resume is None:
                raise Exception("Failed to format resume with GPT")
            self.formatted_resume = formatted_resume
            company_info = self.get_companies_info()
            if company_info:
                formatted_resume['companyBackground'] = company_info

            return {
                "status": "success",
                "uploaded_file": pdf_filename,
                "latest_blob": latest_blob_name,
                "blob_pdf_url": blob_pdf_url,
                "formatted_resume": formatted_resume

            }

        except Exception as e:
            logger.error(f"Error in resume processing workflow: {str(e)}")
            return {
                "status": "error",
                "error_message": str(e),
                "uploaded_file": pdf_filename
            }
